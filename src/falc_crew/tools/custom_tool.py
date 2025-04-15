import os
import json
import re
from crewai.tools import BaseTool
from crewai_tools import RagTool
from typing import List, Optional, Type
from pydantic import BaseModel, Field
from datetime import datetime
from docx import Document
from docx.shared import Pt, Inches


# class MyCustomToolInput(BaseModel):
#     """Input schema for MyCustomTool."""
#     argument: str = Field(..., description="Description of the argument.")

# class MyCustomTool(BaseTool):
#     name: str = "Name of my tool"
#     description: str = (
#         "Clear description for what this tool is useful for, your agent will need this information to use it."
#     )
#     args_schema: Type[BaseModel] = MyCustomToolInput

#     def _run(self, argument: str) -> str:
#         # Implementation goes here
#         return "this is an example of a tool output, ignore it and move along."

# ========== WordExtractorTool ==========
class WordExtractorInput(BaseModel):
    file_path: str = Field(..., description="Chemin vers le document Word source (.docx)")


class WordExtractorTool(BaseTool):
    name: str = "WordExtractorTool"
    description: str = "Lit le contenu texte d'un document Word (.docx) et retourne le texte brut."
    args_schema: Type[BaseModel] = WordExtractorInput

    def _run(self, file_path: str) -> str:
        if not os.path.exists(file_path):
            return "⚠️ Le fichier spécifié est introuvable."

        doc = Document(file_path)
        text = "\n".join([para.text for para in doc.paragraphs if para.text.strip()])
        return text

# ========== FalcDocxStructureTaggerTool ==========

class FalcDocxStructureTaggerInput(BaseModel):
    paragraphs: List[str] = Field(..., description="List of paragraphs extracted from the original .docx file.")

class FalcDocxStructureTaggerTool(BaseTool):
    name: str = "FalcDocxStructureTaggerTool"
    description: str = (
        "Receives a list of paragraphs from the original .docx and tags which paragraph(s) represent the subject and the body."
    )
    args_schema: Type[BaseModel] = FalcDocxStructureTaggerInput

    def _run(self, paragraphs: List[str]) -> str:
        numbered = "\n".join([f"{i}. {p}" for i, p in enumerate(paragraphs)])
        prompt = f"""Here is a list of Word doc paragraphs. Identify:
        - the paragraph index of the subject
        - paragraph indexes of the body

        Paragraphs:
        {numbered}

        Return only JSON like:
        {{ "subject": [index], "body": [i, j, ...] }}"""

        from openai import OpenAI
        client = OpenAI()
        response = client.chat.completions.create(
            model="gpt-4",
            messages=[{"role": "user", "content": prompt}],
        )
        return response.choices[0].message.content


# ========== FalcDocxWriterTool ==========
class FalcDocxWriterInput(BaseModel):
    header: Optional[str] = Field(None, description="Header block for the letter, typically sender info")
    recipient: Optional[str] = Field(None, description="Recipient block, usually address and name")
    subject: Optional[str] = Field(None, description="Subject line of the letter")
    body_sections: List[str] = Field(..., description="List of paragraphs or markdown strings to be rendered")
    footer: Optional[str] = Field(None, description="Final line or sign-off")
    markdown_text: Optional[str] = Field(None, description="Fallback full markdown text")
    original_file: Optional[str] = Field(None, description="Path to the original .docx file")
    subject_index: Optional[int] = Field(None, description="Index of the subject paragraph to replace in the original file")
    body_indexes: Optional[List[int]] = Field(None, description="Indexes of body paragraphs to replace in the original file")
    output_dir: Optional[str] = Field(None, description="Optional directory path where to save the output docx")


class FalcDocxWriterTool(BaseTool):
    name: str = "FalcDocxWriterTool"
    description: str = (
        "Generate a structured FALC Word letter layout with optional header, subject, recipient and footer"
        "This tool aslo scans text for icon placeholders (e.g. [[ICON:key]]) and inserts the corresponding PNG image."
    )
    args_schema: Type[BaseModel] = FalcDocxWriterInput

    def load_icons_map(self) -> dict:
        icons_path = os.path.join("knowledge", "icons.json")
        if not os.path.exists(icons_path):
            return {}
        with open(icons_path, "r", encoding="utf-8") as f:
            return json.load(f)

    def _insert_text_and_icons(self, paragraph, text, icons_map):
        """
        Split the text by icon placeholders and add text runs and image runs.
        The placeholder format is assumed to be [[ICON:KEY]].
        """
        # Regex pattern to detect the placeholder pattern.
        pattern = r'\[\[ICON:(.*?)\]\]'
        # Split the text (keeping the delimiters)
        parts = re.split(f'({pattern})', text)
        for part in parts:
            if part.startswith("[[ICON:") and part.endswith("]]"):
                # Extract the key
                icon_key = re.findall(pattern, part)
                if icon_key:
                    key = icon_key[0].strip()
                    image_path = icons_map.get(key)
                    if image_path and os.path.exists(image_path):
                        run = paragraph.add_run()
                        # Adjust the image size as needed (e.g., width=Inches(0.2))
                        run.add_picture(image_path, width=Inches(0.2))
                    else:
                        # If image is not found, insert the placeholder as text.
                        paragraph.add_run(f"[Missing icon: {key}]")
            else:
                # Remove icon labels like 'direction' after [[ICON:direction]]
                cleaned = part.strip()
                if cleaned in icons_map:
                    # Skip it — already shown as an icon
                    continue

                paragraph.add_run(part)

    def _run(
        self,
        header=None,
        recipient=None,
        subject=None,
        body_sections=None,
        footer=None,
        markdown_text=None,
        original_file=None,
        subject_index=None,
        body_indexes=None,
        output_dir=None
    ) -> str:
        icons_map = self.load_icons_map()

        if original_file and subject_index is not None and body_indexes:
            # 🔁 Rewrite mode
            doc = Document(original_file)
            paragraphs = doc.paragraphs

            if 0 <= subject_index < len(paragraphs):
                para = paragraphs[subject_index]
                para.clear()
                self._insert_text_and_icons(para, subject, icons_map)

            for i, section in zip(body_indexes, body_sections):
                if 0 <= i < len(paragraphs):
                    para = paragraphs[i]
                    para.clear()
                    self._insert_text_and_icons(para, section, icons_map)

            timestamp = datetime.now().strftime("%Y%m%d_%H%M")
            original_name = os.path.splitext(os.path.basename(original_file))[0]
            output_filename = f"{original_name}_falc_{timestamp}.docx"
            output_dir = output_dir or "output"
            os.makedirs(output_dir, exist_ok=True)
            output_path = os.path.join(output_dir, output_filename)

        else:
            # 🆕 Build new layout
            doc = Document()

            # Style
            style = doc.styles['Normal']
            font = style.font
            font.name = 'Arial'
            font.size = Pt(10)

            # Header/Recipient
            if header or recipient:
                table = doc.add_table(rows=1, cols=2)
                table.autofit = False
                hdr_cell, rcpt_cell = table.rows[0].cells
                if header:
                    hdr_cell.text = header
                if recipient:
                    rcpt_cell.text = recipient
                doc.add_paragraph("")

            # Subject
            if subject:
                subject_paragraph = doc.add_paragraph(subject)
                subject_paragraph.style = 'Heading 2'

            # Body
            if body_sections:
                for section in body_sections:
                    clean = section.strip()
                    if not clean:
                        continue
                    if clean.startswith("##"):
                        doc.add_paragraph(clean.replace("##", "").strip(), style='Heading 3')
                    else:
                        p = doc.add_paragraph()
                        self._insert_text_and_icons(p, clean, icons_map)
                        p.paragraph_format.space_after = Pt(10)
                        p.paragraph_format.line_spacing = 1.5
            elif markdown_text:
                for line in markdown_text.splitlines():
                    p = doc.add_paragraph(line)
                    p.paragraph_format.space_after = Pt(10)
                    p.paragraph_format.line_spacing = 1.5

            # Footer
            if footer:
                doc.add_paragraph("\n" + footer)

            output_dir = output_dir or "output"
            os.makedirs(output_dir, exist_ok=True)
            output_path = os.path.join(output_dir, "falc_translated_output.docx")


        os.makedirs("output", exist_ok=True)
        doc.save(output_path)
        return f"✅ FALC document saved: {output_path}"


# ========== FalcIconLookupTool ==========
class FalcIconLookupInput(BaseModel):
    """No input needed for this tool."""
    pass

class FalcIconLookupTool(BaseTool):
    name: str = "FalcIconLookupTool"
    description: str = (
        "Provides a dictionary of available icons based on keyword-to-emoji mapping "
        "from the icons.json file, so the agent can decide where to use them."
    )
    args_schema: Type[BaseModel] = FalcIconLookupInput

    def _run(self) -> str:
        icons_path = os.path.join("knowledge", "icons.json")
        if not os.path.exists(icons_path):
            return "⚠️ Error: icons.json file not found in the 'knowledge/' directory."
        with open(icons_path, "r", encoding="utf-8") as f:
            icon_map = json.load(f)

        # Format the icons list. If the value is a PNG file, assume it needs to be rendered as an image.
        formatted_entries = []
        for key, value in icon_map.items():
            if isinstance(value, str) and value.lower().endswith(".png"):
                # Display a markdown-style image link (adjust based on your UI requirements)
                formatted_entries.append(f"- **{key}**: ![]({value})")
            else:
                formatted_entries.append(f"- **{key}**: {value}")
        formatted = "\n".join(formatted_entries)
        return f"📙 Available company icons:\n{formatted}"



# ========== ReferenceModelRetrieverTool ==========

class ReferenceModelRetrieverInput(BaseModel):
    query: str = Field(..., description="Query to search for a similar FALC translation model.")

class ReferenceModelRetrieverTool(RagTool):
    name: str = "ReferenceModelRetriever"
    description: str = (
        "Use this tool to search a bank of reference FALC translations "
        "and find similar documents based on your input."
    )
    args_schema: Type[BaseModel] = ReferenceModelRetrieverInput
